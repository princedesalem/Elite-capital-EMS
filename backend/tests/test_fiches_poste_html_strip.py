"""Tests de la conversion .docx → HTML mammoth + strip des deux dernières sections."""
from io import BytesIO

import pytest


def _build_docx(blocks):
    """Construit un .docx en mémoire avec python-docx à partir d'une liste de blocs.
    Chaque bloc est soit:
      - {'h': str}   → titre Heading 1
      - {'p': str}   → paragraphe
      - {'table': [[c1,c2,...], ...]} → tableau
    """
    from docx import Document
    doc = Document()
    for b in blocks:
        if 'h' in b:
            doc.add_heading(b['h'], level=1)
        elif 'p' in b:
            doc.add_paragraph(b['p'])
        elif 'table' in b:
            rows = b['table']
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    t.rows[ri].cells[ci].text = val
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def test_docx_to_html_strips_rubrique_and_validation_sections():
    pytest.importorskip('mammoth')
    pytest.importorskip('bs4')
    pytest.importorskip('docx')

    from app.routers.fiches_poste_router import _docx_to_clean_html

    blocks = [
        {'h': 'Mission Principale'},
        {'p': "Le titulaire doit assurer la coordination des activités."},
        {'h': 'Compétences Requises'},
        {'p': 'Leadership, communication, rigueur.'},
        # Rubrique / Employé table to strip
        {'table': [
            ['Rubrique', "Nom de l'employé"],
            ['Signature', '____________'],
            ['Date', '____________'],
        ]},
        {'h': 'Fiche de validation'},
        {'p': 'Validé par DG le …'},
    ]
    docx_bytes = _build_docx(blocks)
    html = _docx_to_clean_html(docx_bytes)

    assert 'Mission Principale' in html
    assert 'coordination' in html
    assert 'Compétences Requises' in html
    # The two last sections must be removed
    lower = html.lower()
    assert 'rubrique' not in lower
    assert "fiche de validation" not in lower
    assert "validé par dg" not in lower


def test_docx_to_html_handles_no_strip_markers():
    """Si aucune section à supprimer, le HTML doit être intact."""
    pytest.importorskip('mammoth')
    pytest.importorskip('docx')
    from app.routers.fiches_poste_router import _docx_to_clean_html

    blocks = [
        {'h': 'Section A'},
        {'p': 'Contenu A.'},
        {'h': 'Section B'},
        {'p': 'Contenu B.'},
    ]
    html = _docx_to_clean_html(_build_docx(blocks))
    assert 'Section A' in html
    assert 'Section B' in html
    assert 'Contenu B' in html
