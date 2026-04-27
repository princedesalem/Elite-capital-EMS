from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Page margins - smaller for mobile reading
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# --- Helper functions ---
def add_title(doc, text, size=22, color=RGBColor(0x1A, 0x37, 0x6C), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_subtitle(doc, text, size=13, color=RGBColor(0x2E, 0x86, 0xAB), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_section_header(doc, text, size=13, color=RGBColor(0x1A, 0x37, 0x6C)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, indent=0, bold_prefix=None, color=RGBColor(0x22, 0x22, 0x22), size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2 + indent * 0.2)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix + " ")
        run_bold.bold = True
        run_bold.font.size = Pt(size)
        run_bold.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E86AB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(size)

# =====================
# PAGE 1 - PRESENTATION
# =====================

add_spacer(doc, 4)
add_title(doc, "EMS — Enterprise Management System", size=20)
add_subtitle(doc, "Elite Capital Group", size=13, bold=True, color=RGBColor(0x2E, 0x86, 0xAB))
add_subtitle(doc, "Plateforme de gestion interne intelligente", size=11, color=RGBColor(0x55, 0x55, 0x55))
add_divider(doc)

# Qu'est-ce que EMS
add_section_header(doc, "📌  Qu'est-ce que EMS ?")
add_bullet(doc, "Une plateforme numérique centralisant la gestion de l'entreprise au quotidien")
add_bullet(doc, "Accessible depuis n'importe quel appareil — ordinateur, tablette ou téléphone")
add_bullet(doc, "Conçue sur mesure pour Elite Capital Group")
add_bullet(doc, "Sécurisée, traçable et évolutive")

add_divider(doc)

# Ce que EMS gère déjà
add_section_header(doc, "✅  Ce qu'EMS gère déjà")
add_bullet(doc, "Ressources humaines :", bold_prefix="RH")
add_bullet(doc, "Congés, permissions, missions, évaluations, remplaçants", indent=1)
add_bullet(doc, "Structure organisationnelle, hiérarchie, localisations", indent=1)
add_bullet(doc, "Workflow de validation à plusieurs niveaux (Directeur → DG → RH)", indent=1)

add_bullet(doc, "Tableau de bord personnalisé par rôle :", bold_prefix="Analytics")
add_bullet(doc, "Chaque collaborateur voit ses propres données, son responsable voit son équipe", indent=1)
add_bullet(doc, "La direction dispose d'une vue globale consolidée", indent=1)

add_bullet(doc, "Sécurité et accès :", bold_prefix="Sécurité")
add_bullet(doc, "Authentification à deux facteurs (2FA)", indent=1)
add_bullet(doc, "Gestion des rôles et permissions par profil", indent=1)

add_divider(doc)

# Ce qui arrive prochainement
add_section_header(doc, "🔜  Modules en cours de déploiement")
add_bullet(doc, "Module Achats — bons de commande, factures fournisseurs")
add_bullet(doc, "Module Commercial — gestion clients, devis, contrats")
add_bullet(doc, "Module CRM — suivi des relations et interactions clients")
add_bullet(doc, "Module Flotte — gestion du parc automobile")
add_bullet(doc, "Module Audit & Projets — suivi des missions d'inspection")
add_bullet(doc, "Module Communication interne")

add_divider(doc)

# === IA ===
add_section_header(doc, "🤖  Intelligence Artificielle intégrée dans EMS", color=RGBColor(0x6A, 0x0D, 0xAD))

add_bullet(doc, "Credit Scoring automatique :", bold_prefix="💳", color=RGBColor(0x6A, 0x0D, 0xAD))
add_bullet(doc, "Évaluation automatique du profil financier de chaque partenaire ou client", indent=1)
add_bullet(doc, "Score calculé en temps réel à partir des données disponibles dans EMS", indent=1)
add_bullet(doc, "Aide à la décision immédiate pour les engagements commerciaux", indent=1)

add_bullet(doc, "Analyse prédictive RH :", bold_prefix="📊", color=RGBColor(0x6A, 0x0D, 0xAD))
add_bullet(doc, "Détection précoce des risques d'absentéisme ou de turnover", indent=1)
add_bullet(doc, "Recommandations automatiques sur les besoins en ressources", indent=1)

add_bullet(doc, "Assistant intelligent :", bold_prefix="💬", color=RGBColor(0x6A, 0x0D, 0xAD))
add_bullet(doc, "Réponses automatiques aux questions courantes des employés (congés, soldes, missions...)", indent=1)
add_bullet(doc, "Résumé automatique des dossiers et rapports", indent=1)

add_bullet(doc, "Alertes et anomalies :", bold_prefix="🔔", color=RGBColor(0x6A, 0x0D, 0xAD))
add_bullet(doc, "Détection automatique d'incohérences dans les données ou les dépenses", indent=1)
add_bullet(doc, "Priorisation intelligente des tâches et demandes en attente", indent=1)

add_bullet(doc, "Rapports narratifs automatiques :", bold_prefix="📄", color=RGBColor(0x6A, 0x0D, 0xAD))
add_bullet(doc, "Génération automatique de rapports en langage naturel à partir des données", indent=1)
add_bullet(doc, "Export PDF ou Word en un clic", indent=1)

add_divider(doc)

# Bénéfices
add_section_header(doc, "🎯  Ce que ça apporte concrètement")
add_bullet(doc, "Gain de temps significatif sur les tâches administratives récurrentes")
add_bullet(doc, "Décisions basées sur des données réelles, pas sur des intuitions")
add_bullet(doc, "Moins d'erreurs humaines — processus automatisés et validés")
add_bullet(doc, "Meilleure traçabilité de toutes les opérations internes")
add_bullet(doc, "Image moderne et professionnelle pour Elite Capital")

add_divider(doc)

# Footer
add_spacer(doc, 4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Elite Capital Group  ·  EMS v1.0  ·  {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.italic = True

# Save
output_path = r"c:\Users\cedri\OneDrive - ELITE CAPITAL Group S.A\Documents\EMS\EMS_Presentation_Direction.docx"
doc.save(output_path)
print(f"Document créé : {output_path}")
