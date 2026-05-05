"""
Génère le document Word des prérequis de déploiement — Elite Capital EMS
Charte graphique : navy #02162e | rouge #d0202b | police Century Gothic
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs charte ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x02, 0x16, 0x2e)
RED    = RGBColor(0xD0, 0x20, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)   # fond ligne paire tableau
GREY   = RGBColor(0x55, 0x65, 0x78)   # texte secondaire
FONT   = "Century Gothic"

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_border(cell, positions=("top","bottom","left","right")):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for pos in positions:
        b = OxmlElement(f"w:{pos}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D0202B")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font=FONT):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1):
    """Titre de section avec filet rouge en bas."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    if level == 1:
        add_run(para, text, bold=True, size=15, color=NAVY)
    else:
        add_run(para, text, bold=True, size=12, color=RED)
    # filet bas rouge
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D0202B")
    pBdr.append(bot)
    pPr.append(pBdr)
    return para

def body(doc, text, indent=False, bold=False, size=10.5, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    add_run(para, text, bold=bold, size=size, color=color or GREY)
    return para

def bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    para.paragraph_format.space_after  = Pt(2)
    add_run(para, text, size=10)
    return para

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # en-tête navy
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "02162E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=10, color=WHITE)

    # données
    for r_idx, row_data in enumerate(rows):
        bg = "F0F4F8" if r_idx % 2 == 0 else "FFFFFF"
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bold_cell = (c_idx == 0)
            col = NAVY if bold_cell else GREY
            add_run(p, str(cell_text), bold=bold_cell, size=9.5, color=col)

    # largeurs colonnes
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])

    doc.add_paragraph()   # espacement après tableau
    return table

# ══════════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # ── Marges ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══ PAGE DE COUVERTURE ════════════════════════════════════════════════════
    # Bandeau navy
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(0)
    cover.paragraph_format.space_after  = Pt(0)
    pPr  = cover._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "02162E")
    pPr.append(shd)
    add_run(cover, "   ", size=4)   # padding

    # Titre principal
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1.paragraph_format.space_before = Pt(30)
    t1.paragraph_format.space_after  = Pt(8)
    add_run(t1, "ELITE CAPITAL", bold=True, size=28, color=NAVY)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(6)
    add_run(t2, "Enterprise Management System", bold=False, size=16, color=RED)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.paragraph_format.space_after = Pt(40)
    add_run(t3, "Prérequis de déploiement local", bold=True, size=20, color=NAVY)

    # Sous-titre
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    add_run(sub, "Spécifications matérielles, logicielles et dépendances", size=12, color=GREY)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(50)
    add_run(date_para, "Mai 2026  •  Version 1.0", size=11, italic=True, color=GREY)

    # Filet rouge de séparation
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(60)
    pPr2 = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D0202B")
    pBdr.append(bot)
    pPr2.append(pBdr)

    doc.add_page_break()

    # ══ INTRODUCTION ══════════════════════════════════════════════════════════
    heading(doc, "1.  Contexte et périmètre")
    body(doc,
         "Ce document recense l'ensemble des prérequis nécessaires au déploiement local "
         "de l'application Elite Capital EMS (Employee Management System). Il couvre "
         "les spécifications matérielles recommandées, les logiciels système, les "
         "outils de développement ainsi que toutes les bibliothèques tierces utilisées "
         "par l'application.")
    body(doc,
         "Les valeurs indiquées intègrent une marge de croissance pour l'intégration "
         "des modules futurs (Achats, Commercial, Marketing) prévus en Phase 4.")

    heading(doc, "Architecture applicative", level=2)
    add_table(doc,
        ["Composant", "Technologie", "Port", "Description"],
        [
            ["Frontend",      "React 18 + Vite 5",       "5173", "Interface utilisateur SPA"],
            ["Backend API",   "FastAPI + Python 3.11",   "8000", "API REST + scheduler APScheduler"],
            ["Base de données","MySQL 8.0",              "3307", "Stockage relationnel principal"],
            ["Reverse Proxy", "Vite dev proxy (dev) / Nginx (prod)", "80/443", "Routage HTTP"],
            ["DevOps Manager","FastAPI local",           "9000", "CI/CD, backup, release pipeline"],
        ],
        col_widths=[3.5, 5.0, 2.2, 6.3]
    )

    # ══ PRÉREQUIS MATÉRIELS ═══════════════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "2.  Prérequis matériels")

    body(doc,
         "Les spécifications ci-dessous sont données pour un déploiement Dockerisé "
         "sur un serveur ou poste de travail dédié. Les valeurs « Recommandé avec "
         "modules futurs » anticipent l'ajout des modules Achats, Commercial, "
         "Marketing et une augmentation du nombre d'utilisateurs simultanés.")

    heading(doc, "2.1  Processeur (CPU)", level=2)
    add_table(doc,
        ["Critère", "Minimum (état actuel)", "Recommandé (modules futurs)"],
        [
            ["Cœurs physiques",    "4 cœurs",             "8 cœurs ou plus"],
            ["Architecture",       "x86-64 (AMD64)",       "x86-64 (AMD64)"],
            ["Fréquence",          "2,0 GHz",              "3,0 GHz ou plus"],
            ["Virtualisation",     "Intel VT-x ou AMD-V activé (requis Docker)", "Intel VT-x ou AMD-V activé"],
            ["Hyperthreading",     "Non obligatoire",      "Recommandé"],
        ],
        col_widths=[5.0, 6.0, 6.0]
    )

    heading(doc, "2.2  Mémoire vive (RAM)", level=2)
    add_table(doc,
        ["Critère", "Minimum (état actuel)", "Recommandé (modules futurs)"],
        [
            ["RAM totale",         "8 Go",                 "16 Go"],
            ["RAM libre pour Docker", "6 Go",              "12 Go"],
            ["Type",               "DDR4 ou supérieur",    "DDR4 2666 MHz ou supérieur"],
            ["Répartition Docker",
             "MySQL 1,5 Go | Backend 1,5 Go | Frontend 0,5 Go | Système 2,5 Go",
             "MySQL 3 Go | Backend 3 Go | Frontend 1 Go | Système 3 Go"],
        ],
        col_widths=[4.5, 6.5, 6.0]
    )

    heading(doc, "2.3  Stockage", level=2)
    add_table(doc,
        ["Critère", "Minimum (état actuel)", "Recommandé (modules futurs)"],
        [
            ["Espace disque total", "100 Go",             "250 Go"],
            ["Type de disque",      "SSD SATA",           "SSD NVMe"],
            ["Répartition",
             "OS + Docker 30 Go | Images Docker 20 Go | BD + uploads 30 Go | Logs/sauvegardes 20 Go",
             "OS + Docker 40 Go | Images 30 Go | BD + uploads 100 Go | Logs/sauvegardes 80 Go"],
            ["IOPS lecture/écriture", "500 / 300 IOPS",  "3 000 / 1 500 IOPS"],
            ["Sauvegarde",          "Support externe recommandé", "NAS ou cloud (sauvegarde quotidienne)"],
        ],
        col_widths=[4.5, 6.5, 6.0]
    )

    heading(doc, "2.4  Réseau", level=2)
    add_table(doc,
        ["Critère", "Minimum", "Recommandé"],
        [
            ["Bande passante LAN",   "100 Mbps",   "1 Gbps"],
            ["Interface réseau",     "1 carte Ethernet", "1 carte Ethernet Gigabit"],
            ["Connectivité Internet","Requise (pull images Docker, mises à jour)", "Fibre ou ADSL stable ≥ 10 Mbps"],
            ["Ports à ouvrir (LAN)", "5173 (front), 8000 (API), 3307 (MySQL)", "Idem + 9000 (DevOps Manager)"],
            ["DNS local",            "Non obligatoire",  "Recommandé pour nom de domaine interne"],
        ],
        col_widths=[4.5, 5.5, 7.0]
    )

    heading(doc, "2.5  Système d'exploitation", level=2)
    add_table(doc,
        ["OS", "Version minimale", "Notes"],
        [
            ["Windows",    "Windows 10 Pro / Server 2019",
             "WSL2 requis pour Docker Desktop. Hyper-V ou WSL2 activé."],
            ["Ubuntu Linux", "Ubuntu 22.04 LTS (Jammy)",
             "Recommandé pour serveur de production. Docker Engine natif."],
            ["Debian",     "Debian 11 (Bullseye) ou supérieur",
             "Docker Engine natif. Léger pour serveur dédié."],
            ["macOS",      "macOS 12 Monterey ou supérieur",
             "Apple Silicon (M1/M2) compatible Docker Desktop."],
        ],
        col_widths=[3.0, 5.0, 9.0]
    )

    # ══ PRÉREQUIS LOGICIELS ═══════════════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "3.  Applications et outils système à installer")

    body(doc,
         "L'application tourne intégralement dans Docker Compose. Les outils "
         "suivants doivent être installés sur le serveur hôte.", bold=False)

    heading(doc, "3.1  Outils obligatoires", level=2)
    add_table(doc,
        ["Application", "Version minimale", "Rôle", "Lien officiel"],
        [
            ["Docker Desktop\n(Windows/macOS)\nou Docker Engine\n(Linux)",
             "Docker 24.x\nCompose v2.20+",
             "Conteneurisation de toute la stack (backend, frontend, MySQL)",
             "https://docs.docker.com/get-docker/"],
            ["Docker Compose",
             "v2.20 (plugin intégré à Docker Desktop)",
             "Orchestration multi-conteneurs via docker-compose.yml",
             "Inclus dans Docker Desktop"],
            ["Git",
             "2.40 ou supérieur",
             "Clonage du dépôt, gestion des versions, déploiement",
             "https://git-scm.com/downloads"],
            ["Python",
             "3.11.x",
             "Scripts de migration DB, scripts utilitaires hors Docker",
             "https://www.python.org/downloads/"],
            ["Node.js (LTS)",
             "18.x LTS ou 20.x LTS",
             "Exécution Vite / Vitest / Playwright hors Docker (optionnel en dev)",
             "https://nodejs.org/en/download/"],
        ],
        col_widths=[3.8, 3.2, 6.0, 4.0]
    )

    heading(doc, "3.2  Outils recommandés (développement et administration)", level=2)
    add_table(doc,
        ["Application", "Version", "Rôle", "Lien"],
        [
            ["Visual Studio Code",
             "Dernière stable",
             "Éditeur principal — extensions Docker, Python, ESLint, Pylance",
             "https://code.visualstudio.com/"],
            ["MySQL Workbench",
             "8.0 ou supérieur",
             "Administration visuelle de la base de données MySQL",
             "https://dev.mysql.com/downloads/workbench/"],
            ["Postman / Insomnia",
             "Dernière stable",
             "Test et documentation des endpoints API REST",
             "https://www.postman.com/"],
            ["Windows Subsystem for Linux 2 (WSL2)",
             "WSL2 (Windows uniquement)",
             "Requis par Docker Desktop sur Windows pour les performances",
             "https://learn.microsoft.com/fr-fr/windows/wsl/"],
            ["PowerShell",
             "5.1 (inclus Windows) ou 7.x",
             "Exécution des scripts .ps1 (backup-db, restore-db, run-dev)",
             "Inclus dans Windows"],
            ["curl / wget",
             "Tout récent",
             "Vérification endpoints API, téléchargements",
             "Inclus dans Linux/macOS, winget sur Windows"],
        ],
        col_widths=[3.8, 3.2, 6.0, 4.0]
    )

    heading(doc, "3.3  Ports système requis (à libérer/ouvrir)", level=2)
    add_table(doc,
        ["Port", "Protocole", "Service", "Obligatoire"],
        [
            ["5173", "TCP", "Frontend React (Vite dev server)", "Oui"],
            ["8000", "TCP", "Backend FastAPI (API REST + WebSocket)", "Oui"],
            ["3307", "TCP", "MySQL 8.0 (host → conteneur 3306)", "Oui"],
            ["9000", "TCP", "DevOps Manager (CI/CD, backup, health)", "Recommandé"],
            ["587",  "TCP", "SMTP sortant Gmail (envoi emails)", "Si SMTP_ENABLED=true"],
            ["443",  "TCP", "HTTPS sortant (GitHub API, WebPush)", "Oui"],
        ],
        col_widths=[2.0, 3.0, 7.0, 5.0]
    )

    # ══ DÉPENDANCES PYTHON ════════════════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "4.  Bibliothèques Python (Backend)")

    body(doc,
         "Installées automatiquement via pip dans le conteneur Docker backend "
         "(voir backend/requirements.txt). Aucune installation manuelle requise "
         "si Docker est utilisé.")

    heading(doc, "4.1  Framework et API", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["fastapi",           "≥ 0.110",  "Framework web asynchrone — cœur de l'API REST"],
            ["uvicorn[standard]", "≥ 0.29",   "Serveur ASGI pour FastAPI (avec rechargement automatique)"],
            ["pydantic",          "v2.x",     "Validation et sérialisation des schémas de données"],
            ["python-multipart",  "≥ 0.0.9",  "Upload de fichiers (photos employés, imports CSV/XLSX)"],
            ["python-dotenv",     "≥ 1.0",    "Lecture des variables d'environnement (.env)"],
        ],
        col_widths=[4.0, 3.0, 10.0]
    )

    heading(doc, "4.2  Base de données et ORM", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["SQLAlchemy",     "≥ 2.0",   "ORM — mapping objet-relationnel avec MySQL"],
            ["alembic",        "≥ 1.13",  "Migrations de schéma de base de données"],
            ["pymysql",        "≥ 1.1",   "Driver Python pur pour connexion MySQL 8.0"],
            ["psycopg2-binary","≥ 2.9",   "Driver PostgreSQL (réservé pour migrations futures)"],
        ],
        col_widths=[4.0, 3.0, 10.0]
    )

    heading(doc, "4.3  Authentification et sécurité", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["python-jose[cryptography]", "≥ 3.3", "Génération et vérification des tokens JWT"],
            ["passlib[bcrypt]",           "≥ 1.7", "Hachage des mots de passe (bcrypt)"],
            ["pyotp",                     "≥ 2.9",  "Authentification TOTP (2FA/MFA)"],
            ["email-validator",           "≥ 2.1",  "Validation syntaxique des adresses email"],
        ],
        col_widths=[4.5, 2.5, 10.0]
    )

    heading(doc, "4.4  Planification et notifications", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["apscheduler",  "≥ 3.10", "Scheduler de tâches (relances, alertes congés, clôtures auto)"],
            ["requests",     "≥ 2.31", "Requêtes HTTP sortantes (geo-services, webhooks)"],
            ["pywebpush",    "≥ 2.0",  "Notifications push navigateur via protocole VAPID/WebPush"],
        ],
        col_widths=[4.0, 3.0, 10.0]
    )

    heading(doc, "4.5  Import / Export et génération de documents", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["pandas",         "≥ 2.1",  "Traitement et analyse des fichiers tabulaires (CSV, Excel)"],
            ["openpyxl",       "≥ 3.1",  "Lecture/écriture de fichiers .xlsx"],
            ["xlrd",           "≥ 2.0",  "Lecture de fichiers .xls (anciens formats Excel)"],
            ["xlwt",           "≥ 1.3",  "Écriture de fichiers .xls"],
            ["pyodbc",         "≥ 4.0.35 < 6", "Connexion aux bases Access (.mdb/.accdb) via ODBC"],
            ["fpdf2",          "≥ 2.7.9", "Génération de rapports PDF (police Century Gothic / Gothic)"],
            ["python-docx",    "≥ 1.1",  "Génération de documents Word (.docx)"],
            ["mammoth",        "≥ 1.6",  "Conversion Word → HTML pour aperçu"],
            ["beautifulsoup4", "≥ 4.12", "Parsing HTML (nettoyage contenu importé)"],
            ["weasyprint",     "≥ 60.0", "Rendu HTML/CSS vers PDF (fiches, rapports)"],
            ["Pillow",         "≥ 10.0", "Traitement d'images (photos employés, redimensionnement)"],
        ],
        col_widths=[4.0, 3.5, 9.5]
    )

    heading(doc, "4.6  Tests et développement", level=2)
    add_table(doc,
        ["Bibliothèque", "Version", "Rôle"],
        [
            ["pytest",         "≥ 9.0",  "Framework de tests unitaires et d'intégration"],
            ["httpx",          "≥ 0.27", "Client HTTP asynchrone pour les tests FastAPI (TestClient)"],
        ],
        col_widths=[4.0, 3.0, 10.0]
    )

    # ══ DÉPENDANCES NODE.JS ═══════════════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "5.  Bibliothèques Node.js (Frontend)")

    body(doc,
         "Installées automatiquement via npm install dans le conteneur Docker frontend "
         "(voir frontend/package.json). Aucune installation manuelle requise si Docker est utilisé.")

    heading(doc, "5.1  Dépendances de production", level=2)
    add_table(doc,
        ["Package", "Version", "Rôle"],
        [
            ["react",           "^18.2.0",  "Bibliothèque UI principale — composants, hooks, rendu"],
            ["react-dom",       "^18.2.0",  "Rendu React dans le DOM du navigateur"],
            ["react-router-dom","^6.14.1",  "Routage SPA côté client (pages, navigation, guards)"],
            ["axios",           "^1.4.0",   "Client HTTP pour les appels API REST backend"],
            ["dayjs",           "^1.11.9",  "Manipulation et formatage des dates (léger, ~2 ko)"],
            ["jwt-decode",      "^3.1.2",   "Décodage des tokens JWT (lecture payload côté client)"],
            ["lucide-react",    "^0.577.0", "Bibliothèque d'icônes SVG (25+ icônes dans l'app)"],
            ["recharts",        "^3.8.0",   "Graphiques et visualisations (dashboard, analytics)"],
            ["html-to-image",   "^1.11.11", "Export de composants React en image PNG/JPEG"],
            ["jspdf",           "^4.2.1",   "Génération de PDF côté client (rapports, exports)"],
            ["xlsx",            "^0.18.5",  "Lecture/écriture de fichiers Excel côté client"],
        ],
        col_widths=[4.0, 3.0, 10.0]
    )

    heading(doc, "5.2  Dépendances de développement et tests", level=2)
    add_table(doc,
        ["Package", "Version", "Rôle"],
        [
            ["vite",                        "^5.0.0",  "Bundler ultra-rapide (HMR, build, proxy dev)"],
            ["@vitejs/plugin-react",        "^4.0.0",  "Plugin Vite — support JSX et Fast Refresh"],
            ["vitest",                      "^3.2.4",  "Framework de tests unitaires (API compatible Jest)"],
            ["@vitest/ui",                  "^3.2.4",  "Interface graphique pour les résultats Vitest"],
            ["@testing-library/react",      "^16.3.2", "Utilitaires de test — rendu et interaction composants"],
            ["@testing-library/jest-dom",   "^6.9.1",  "Matchers supplémentaires pour assertions DOM"],
            ["jsdom",                       "^26.1.0", "Simulation de DOM navigateur dans Node.js"],
            ["@playwright/test",            "^1.53.0", "Tests E2E — navigation, formulaires, flux complets"],
        ],
        col_widths=[4.5, 3.0, 9.5]
    )

    # ══ DÉPENDANCES SYSTÈME (DANS DOCKER) ═════════════════════════════════════
    doc.add_page_break()
    heading(doc, "6.  Dépendances système Linux (dans le conteneur Docker backend)")

    body(doc,
         "Ces paquets apt sont installés automatiquement lors du build de l'image Docker "
         "backend (Dockerfile). Aucune action manuelle requise.")

    add_table(doc,
        ["Paquet apt", "Rôle"],
        [
            ["build-essential, gcc",                        "Compilateur C/C++ pour les extensions Python natives"],
            ["libpq-dev",                                   "Headers PostgreSQL (psycopg2-binary)"],
            ["unixodbc, unixodbc-dev",                      "Driver ODBC pour connexions Access/DSN (pyodbc)"],
            ["mdbtools",                                    "Lecture des bases Microsoft Access (.mdb/.accdb)"],
            ["libjpeg62-turbo-dev, libpng-dev, zlib1g-dev", "Encodage/décodage images (Pillow, WeasyPrint)"],
            ["libpango-1.0-0, libpangoft2-1.0-0",          "Moteur de rendu texte (WeasyPrint → PDF)"],
            ["libharfbuzz0b",                               "Shaping de polices (WeasyPrint)"],
            ["libfontconfig1",                              "Configuration polices système (WeasyPrint)"],
            ["libcairo2",                                   "Rendu vectoriel 2D (WeasyPrint)"],
            ["libgdk-pixbuf-2.0-0",                        "Traitement images GDK (WeasyPrint)"],
        ],
        col_widths=[6.0, 11.0]
    )

    # ══ VARIABLES D'ENVIRONNEMENT ════════════════════════════════════════════
    heading(doc, "7.  Variables d'environnement à configurer")

    body(doc,
         "Ces variables sont définies dans docker-compose.yml. En production, "
         "elles doivent être externalisées dans un fichier .env sécurisé et non "
         "versionné (ajouté au .gitignore).")

    add_table(doc,
        ["Variable", "Obligatoire", "Valeur par défaut / Exemple", "Description"],
        [
            ["DATABASE_URL",          "Oui", "mysql+pymysql://user:pass@db:3306/EMS_DB",
             "URL de connexion MySQL"],
            ["SECRET_KEY",            "Oui", "⚠ À remplacer en production",
             "Clé de signature JWT — longue chaîne aléatoire"],
            ["ACCESS_TOKEN_EXPIRE_MINUTES", "Non", "60",
             "Durée de vie des tokens JWT en minutes"],
            ["INIT_ADMIN_PW",         "Oui", "ChangeMe123!@#",
             "Mot de passe initial du compte admin (à changer immédiatement)"],
            ["WEBPUSH_PUBLIC_KEY",    "Oui", "Clé VAPID publique",
             "Notifications push navigateur (WebPush)"],
            ["WEBPUSH_PRIVATE_KEY",   "Oui", "Clé VAPID privée",
             "Notifications push navigateur (WebPush)"],
            ["WEBPUSH_SUBJECT",       "Oui", "mailto:admin@elitecapital.com",
             "Contact VAPID obligatoire"],
            ["SMTP_ENABLED",          "Non", "false",
             "Activer l'envoi d'emails (true/false)"],
            ["SMTP_HOST",             "Si SMTP", "smtp.gmail.com",
             "Serveur SMTP sortant"],
            ["SMTP_PORT",             "Si SMTP", "587",
             "Port SMTP (TLS)"],
            ["SMTP_USER",             "Si SMTP", "adresse@gmail.com",
             "Compte expéditeur"],
            ["SMTP_PASS",             "Si SMTP", "App password Gmail",
             "Mot de passe d'application (pas le mdp principal)"],
            ["APP_URL",               "Oui", "http://localhost:5173",
             "URL publique du frontend (pour les liens dans les emails)"],
        ],
        col_widths=[4.5, 2.5, 5.0, 5.0]
    )

    # ══ ANTICIPATION MODULES FUTURS ═══════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "8.  Anticipation — Modules futurs (Phase 4)")

    body(doc,
         "Les modules Achats, Commercial et Marketing seront ajoutés lors de la Phase 4. "
         "Les prérequis supplémentaires attendus sont listés ci-dessous.")

    add_table(doc,
        ["Module futur", "Technologie additionnelle prévisible", "Impact serveur"],
        [
            ["Achats",      "Génération PDF bons de commande, envoi email fournisseurs",
             "RAM +1 Go | Stockage +20 Go (PJ fournisseurs)"],
            ["Commercial",  "CRM léger, gestion clients, suivi devis",
             "RAM +1 Go | Nouvelles tables DB | API externe optionnelle"],
            ["Marketing",   "Gestion campagnes, événements, envois email en masse",
             "RAM +1 Go | Service file d'attente (Celery/Redis) recommandé"],
            ["Reporting avancé", "Génération rapports complexes, graphiques PDF",
             "RAM +500 Mo | Librairie matplotlib ou reportlab"],
            ["Déploiement production", "Reverse proxy Nginx, certificat SSL/TLS, CI/CD GitHub Actions",
             "CPU +2 cœurs | RAM +2 Go | Stockage +30 Go | Port 443"],
        ],
        col_widths=[3.5, 7.5, 6.0]
    )

    heading(doc, "Bibliothèques Python supplémentaires envisagées", level=2)
    add_table(doc,
        ["Bibliothèque", "Usage prévu"],
        [
            ["celery + redis",  "File d'attente asynchrone pour envois email en masse (Marketing)"],
            ["matplotlib",      "Graphiques statistiques export PDF/PNG (Reporting)"],
            ["reportlab",       "Génération PDF avancée (bons de commande, contrats)"],
            ["stripe / payplug","Intégration paiement si module Commercial évolue"],
            ["pycountry",       "Normalisation pays/devises (Commercial international)"],
        ],
        col_widths=[4.0, 13.0]
    )

    # ══ RÉCAPITULATIF ════════════════════════════════════════════════════════
    doc.add_page_break()
    heading(doc, "9.  Récapitulatif — Checklist d'installation")

    body(doc, "Cochez chaque point avant de lancer l'application :", bold=True, color=NAVY)
    doc.add_paragraph()

    checklist = [
        ("Système d'exploitation compatible installé (Windows 10 Pro+ / Ubuntu 22.04+ / Debian 11+)", True),
        ("WSL2 activé (Windows uniquement) + Virtualisation CPU activée dans le BIOS", True),
        ("Docker Desktop (Windows/macOS) ou Docker Engine + Docker Compose v2 (Linux) installé", True),
        ("Git 2.40+ installé et configuré (git config user.name / user.email)", True),
        ("Dépôt cloné : git clone https://github.com/princedesalem/Elite-capital-EMS.git", True),
        ("Python 3.11 installé (pour scripts de migration hors Docker)", False),
        ("Node.js 18 LTS installé (pour développement hors Docker)", False),
        ("Ports 5173, 8000, 3307 libres sur le serveur hôte", True),
        ("Variables d'environnement SECRET_KEY et INIT_ADMIN_PW modifiées", True),
        ("SMTP configuré si envoi d'emails activé (SMTP_ENABLED=true)", False),
        ("Lancer : docker compose up --build -d (première installation)", True),
        ("Vérifier la santé : docker compose ps — tous les services «running»", True),
        ("Changer le mot de passe admin immédiatement après la première connexion", True),
    ]

    for text, required in checklist:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(3)
        marker = "☐  "
        add_run(para, marker, bold=True, size=11, color=RED if required else GREY)
        label = "[OBLIGATOIRE]  " if required else "[OPTIONNEL]  "
        add_run(para, label, bold=True, size=9.5,
                color=RED if required else GREY)
        add_run(para, text, size=10, color=GREY)

    # ══ PIED DE PAGE ═════════════════════════════════════════════════════════
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr3 = footer_para._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"),   "single")
    top2.set(qn("w:sz"),    "6")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "D0202B")
    pBdr2.append(top2)
    pPr3.append(pBdr2)
    add_run(footer_para,
            "Elite Capital Group S.A  •  Document confidentiel  •  Mai 2026",
            size=9, italic=True, color=GREY)

    # ── Sauvegarde ─────────────────────────────────────────────────────────
    out = "PREREQUIS_DEPLOIEMENT_EMS.docx"
    doc.save(out)
    print(f"✅  Document généré : {out}")

if __name__ == "__main__":
    build_document()
