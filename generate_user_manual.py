# -*- coding: utf-8 -*-
"""
Génère MANUEL_UTILISATION_EMS.docx
Manuel d'utilisation complet de l'application EMS — Elite Capital Group
Police : Century Gothic | Charte graphique EMS
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Charte EMS ────────────────────────────────────────────────────────────
BLEU        = RGBColor(0x02, 0x16, 0x2E)   # #02162E
BLEU_CLAIR  = RGBColor(0x17, 0x28, 0x3D)   # #17283D
ROUGE       = RGBColor(0xD0, 0x20, 0x2B)   # #D0202B
GRIS        = RGBColor(0x60, 0x60, 0x60)   # #606060
BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CLAIR  = RGBColor(0xF4, 0xF5, 0xF7)   # #F4F5F7
VERT        = RGBColor(0x2E, 0x7D, 0x32)   # vert validation
FONT        = "Century Gothic"

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D0D5DD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def paragraph_style(para, font=FONT, size=10, color=GRIS, bold=False,
                    italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    for run in para.runs:
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic

def add_run_styled(para, text, font=FONT, size=10, color=GRIS,
                   bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return run

def add_para(doc, text="", font=FONT, size=10, color=GRIS, bold=False,
             italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4,
             space_before=0, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        add_run_styled(p, text, font, size, color, bold, italic)
    return p

def add_heading(doc, text, level=1):
    if level == 1:
        # Grand titre de section (fond bleu)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        shading_elm = OxmlElement("w:pPr")
        run = p.add_run(f"  {text}")
        run.font.name = FONT
        run.font.size = Pt(14)
        run.font.color.rgb = BLANC
        run.font.bold = True
        # Fond bleu via pPr shd
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "02162E")
        pPr.append(shd)
        p.paragraph_format.left_indent = Cm(0)
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"  {text}")
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.color.rgb = BLANC
        run.font.bold = True
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "17283D")
        pPr.append(shd)
        return p
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.color.rgb = ROUGE
        run.font.bold = True
        run.font.underline = True
        return p
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"▸  {text}")
        run.font.name = FONT
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLEU
        run.font.bold = True
        return p

def add_bullet(doc, text, level=0, marker="•"):
    indent = 0.5 + level * 0.5
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    add_run_styled(p, f"{marker}  {text}", size=10, color=GRIS)
    return p

def add_info_box(doc, text, bg="F4F5F7", color=GRIS):
    """Encadré info sur fond gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    run = p.add_run(f"  ℹ  {text}")
    run.font.name = FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(int(bg[:2], 16) ^ 0xFF, 0x28, 0x3D) if bg != "F4F5F7" else BLEU
    run.font.italic = True
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Tableau avec en-tête bleu + lignes alternées."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # En-tête
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "02162E")
        set_cell_borders(cell, "02162E")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(hdr)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BLANC

    # Données
    for ri, row_data in enumerate(rows):
        bg = "F4F5F7" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "D0D5DD")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            run = p.add_run(str(cell_text))
            run.font.name = FONT
            run.font.size = Pt(9)
            run.font.color.rgb = GRIS

    # Largeurs colonnes
    if col_widths:
        for ri2 in range(len(tbl.rows)):
            for ci2, w in enumerate(col_widths):
                tbl.rows[ri2].cells[ci2].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 95)
    run.font.name = FONT
    run.font.size = Pt(5)
    run.font.color.rgb = RGBColor(0xD0, 0xD5, 0xDD)

# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════
doc = Document()

# ── Paramètres page ───────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── Styles de base ────────────────────────────────────────────────────────
for style_name in ("Normal", "Body Text"):
    try:
        s = doc.styles[style_name]
        s.font.name = FONT
        s.font.size = Pt(10)
    except Exception:
        pass

# ═══════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════
# Fond bleu simulé par paragraphes colorés
for _ in range(4):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "02162E")
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

# Logo texte
p_logo = doc.add_paragraph()
pPr = p_logo._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
pPr.append(shd)
p_logo.paragraph_format.space_after = Pt(0)
r = p_logo.add_run("  ELITE CAPITAL GROUP")
r.font.name = FONT; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0xD0, 0x20, 0x2B)

for _ in range(3):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
    pPr.append(shd); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

# Titre principal
p_title = doc.add_paragraph()
pPr = p_title._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
pPr.append(shd)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("MANUEL D'UTILISATION")
r.font.name = FONT; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = BLANC

p_sub = doc.add_paragraph()
pPr = p_sub._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
pPr.append(shd)
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Application EMS")
r.font.name = FONT; r.font.size = Pt(18); r.font.bold = False; r.font.color.rgb = RGBColor(0xD0, 0x20, 0x2B)

p_sub2 = doc.add_paragraph()
pPr = p_sub2._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
pPr.append(shd)
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub2.add_run("Plateforme de Gestion des Ressources Humaines")
r.font.name = FONT; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xB0, 0xC0, 0xD0)

for _ in range(5):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
    pPr.append(shd); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

# Ligne rouge séparatrice
p_sep = doc.add_paragraph()
pPr = p_sep._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D0202B")
pPr.append(shd); p_sep.paragraph_format.space_after = Pt(0); p_sep.paragraph_format.space_before = Pt(0)
p_sep.add_run("  ")

for _ in range(2):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
    pPr.append(shd); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

# Infos bas de page de garde
p_info = doc.add_paragraph()
pPr = p_info._p.get_or_add_pPr()
shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
pPr.append(shd)
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_info.add_run("Version 1.0  ·  Avril 2026  ·  Confidentiel — Usage interne")
r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x90, 0xA0); r.font.italic = True

for _ in range(6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "02162E")
    pPr.append(shd); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

# Saut de page
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "SOMMAIRE", 1)
add_para(doc, "", space_after=4)
sommaire = [
    ("1.", "Accès et Authentification", "3"),
    ("2.", "Rôles et Permissions", "4"),
    ("3.", "Navigation Globale", "5"),
    ("4.", "Page d'Accueil — Dashboard Personnel", "6"),
    ("5.", "Gestion des Employés", "7"),
    ("6.", "Fiche Employé — Création et Modification", "9"),
    ("7.", "Administration (Entités, Directions, Départements, Fonctions)", "11"),
    ("8.", "Gestion des Utilisateurs (Admin)", "12"),
    ("9.", "Congés", "13"),
    ("10.", "Permissions", "15"),
    ("11.", "Demandes de Sorties Professionnelles", "17"),
    ("12.", "Missions", "18"),
    ("13.", "Frais de Mission", "21"),
    ("14.", "Opérations / Workflow Unifié", "22"),
    ("15.", "Remplaçants", "24"),
    ("16.", "Calendrier des Congés", "24"),
    ("17.", "Notifications", "25"),
    ("18.", "Évaluations", "26"),
    ("19.", "Fiches de Poste", "27"),
    ("20.", "Organigramme Interactif", "27"),
    ("21.", "Organisation & Localisation", "28"),
    ("22.", "Profil Personnel et Sécurité", "29"),
    ("23.", "Analytics RH Avancé", "30"),
    ("24.", "Parcours Employé (Timeline)", "30"),
    ("25.", "Import & Export des Employés", "31"),
    ("26.", "Règles Métier Clés", "32"),
    ("27.", "Sécurité et Contrôle d'Accès", "33"),
]
for num, titre, page in sommaire:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    add_run_styled(p, f"{num}  ", size=10, bold=True, color=ROUGE)
    add_run_styled(p, titre, size=10, color=BLEU)
    # Points de suite
    add_run_styled(p, f"  {'·' * max(1, 72 - len(titre))}  {page}", size=9, color=GRIS, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. ACCÈS ET AUTHENTIFICATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  ACCÈS ET AUTHENTIFICATION", 1)
add_para(doc, "L'application EMS est accessible via votre navigateur web à l'adresse communiquée par votre administrateur. Aucune installation n'est requise.", size=10, color=GRIS, space_after=6)

add_heading(doc, "1.1  Connexion", 2)
add_para(doc, "Deux méthodes de connexion sont disponibles :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Matricule + Mot de passe : saisissez votre matricule alphanumérique et votre mot de passe personnel.")
add_bullet(doc, "Lien magique par email : saisissez votre adresse email enregistrée, un lien de connexion sécurisé vous est envoyé.")
add_para(doc, space_after=4)
add_info_box(doc, "Si vous vous connectez pour la première fois, utilisez le mot de passe temporaire transmis par votre administrateur. Vous serez invité à le modifier immédiatement.", bg="FFF0F0")

add_heading(doc, "1.2  Authentification Multi-Facteurs (MFA)", 2)
add_para(doc, "Si votre compte est configuré avec le MFA, un second facteur vous sera demandé après la saisie du mot de passe :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Ouvrez votre application d'authentification (Google Authenticator, Authy, etc.).")
add_bullet(doc, "Saisissez le code à 6 chiffres affiché.")
add_bullet(doc, "Le code est valide 30 secondes. En cas d'expiration, attendez le prochain code.")
add_para(doc, space_after=4)
add_info_box(doc, "L'activation du MFA se fait depuis votre Profil → section Sécurité. L'administrateur peut également l'activer/désactiver depuis la gestion des utilisateurs.")

add_heading(doc, "1.3  Déconnexion", 2)
add_para(doc, "Cliquez sur votre avatar en haut à droite de l'écran, puis sélectionnez « Déconnexion ». Votre session est immédiatement invalidée.", size=10, color=GRIS, space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. RÔLES ET PERMISSIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  RÔLES ET PERMISSIONS", 1)
add_para(doc, "L'accès aux fonctionnalités est déterminé par le rôle attribué à votre compte. Voici la hiérarchie des rôles et leurs droits principaux :", size=10, color=GRIS, space_after=6)

add_table(doc,
    ["Rôle", "Description", "Droits principaux"],
    [
        ["EMPLOYE", "Employé standard", "Voir ses données · Soumettre congés/permissions/missions · Consulter notifications"],
        ["RESPONSABLE", "Chef d'équipe / N+1", "EMPLOYE + Valider congés/missions de son équipe · Analytics département"],
        ["DIRECTEUR", "Chef de direction", "RESPONSABLE + Valider au niveau direction · Analytics direction complète"],
        ["DG", "Directeur Général", "DIRECTEUR + Administration structure · Analytics globale entité"],
        ["RH", "Ressources Humaines", "Accès complet RH · CRUD Employés · Import/Export · Gestion fonctions · Voir tous les workflows"],
        ["DFC", "Directeur Financier et Comptable", "Accès financier · Validation frais · Analytics paie"],
        ["PCA", "Président du Conseil d'Administration", "Accès très élevé · Administration complète · Validation stratégique"],
        ["AG", "Assemblée Générale", "Même niveau PCA · Validation opérations critiques"],
        ["ADMIN", "Administrateur Système", "Accès total · Gestion utilisateurs · Audit logs · Statistiques usage · CI/CD"],
    ],
    col_widths=[3.2, 4.2, 9.5]
)

add_heading(doc, "2.1  Données visibles selon le rôle", 2)
add_para(doc, "Chaque rôle voit les données dans son périmètre hiérarchique :", size=10, color=GRIS, space_after=4)
add_table(doc,
    ["Rôle", "Périmètre employés visibles", "Salaires visibles"],
    [
        ["EMPLOYE", "Soi-même uniquement", "Son propre salaire"],
        ["RESPONSABLE", "Son équipe directe", "Non"],
        ["DIRECTEUR", "Sa direction entière", "Non"],
        ["DG / PCA / AG", "Toute l'entité", "Oui (tous)"],
        ["RH", "Tous les employés", "Oui (tous)"],
        ["ADMIN", "Tous les employés", "Oui (tous)"],
    ],
    col_widths=[3.5, 7.5, 5.9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. NAVIGATION GLOBALE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  NAVIGATION GLOBALE", 1)

add_heading(doc, "3.1  Barre supérieure (Navbar)", 2)
add_para(doc, "La barre supérieure est visible sur toutes les pages et comprend :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Logo Elite Capital (lien vers l'accueil)")
add_bullet(doc, "Icône de cloche 🔔 : accès au centre de notifications (badge rouge indique les non lues)")
add_bullet(doc, "Avatar utilisateur avec votre prénom/nom : cliquez pour accéder à votre profil ou vous déconnecter")

add_heading(doc, "3.2  Sidebar (Menu latéral gauche)", 2)
add_para(doc, "Le menu latéral regroupe tous les modules de l'application. Il peut être réduit en cliquant sur l'icône de hamburger. Les modules visibles dépendent de votre rôle :", size=10, color=GRIS, space_after=4)

add_table(doc,
    ["Module", "Sous-modules / Items", "Rôles requis"],
    [
        ["Accueil", "Dashboard Personnel", "Tous"],
        ["Ressources Humaines", "Employés · Administration · Absences · Missions · Évaluations · Fiches de poste · Calendrier · Workflow · Événements · Analytics · Parcours · Performances 360 · Planification · Talents · Club Review", "Tous (selon rôle)"],
        ["Achats / Commercial / Marketing", "Placeholder (bientôt disponible)", "Tous (lecture)"],
        ["SI / Flotte / Audit / Projets / CRM", "Placeholder (bientôt disponible)", "Tous (lecture)"],
    ],
    col_widths=[3.5, 9, 4.4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. PAGE D'ACCUEIL
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "4.  PAGE D'ACCUEIL — DASHBOARD PERSONNEL", 1)
add_para(doc, "La page d'accueil (/rh/home) est le premier écran après connexion. Elle offre une vue personnalisée de votre activité et de celle de votre équipe.", size=10, color=GRIS, space_after=6)

add_heading(doc, "4.1  Profil résumé", 2)
add_para(doc, "En haut de la page, un encadré affiche votre photo, nom, prénom, matricule, rôle, fonction et email. Un lien « Modifier profil » permet d'accéder directement à vos informations.", size=10, color=GRIS, space_after=4)

add_heading(doc, "4.2  KPIs rapides", 2)
add_bullet(doc, "Total des employés dans votre périmètre")
add_bullet(doc, "Nombre de congés en attente de validation")
add_bullet(doc, "Opérations en cours vous concernant")
add_para(doc, space_after=4)

add_heading(doc, "4.3  Team Engagement Space", 2)
add_para(doc, "Espace social collaboratif permettant d'animer la vie d'équipe :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Shoutouts : messages d'encouragement publics à un collègue")
add_bullet(doc, "Kudos : reconnaissance de valeurs (Excellence, Innovation, etc.)")
add_bullet(doc, "Sondages (Polls) : créer des sondages rapides pour l'équipe")
add_bullet(doc, "Vœux (Wishes) : envoyer des vœux d'anniversaire ou de départ")
add_bullet(doc, "Flux Opérations : voir les demandes récentes en temps réel")
add_para(doc, space_after=4)

add_heading(doc, "4.4  Actions rapides", 2)
add_bullet(doc, "Bouton « Demander Congé » → ouvre directement le formulaire de congé")
add_bullet(doc, "Bouton « Nouvelle Mission » → ouvre le formulaire de mission")
add_bullet(doc, "Bouton « Demander Permission » → ouvre le formulaire de permission")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. GESTION DES EMPLOYÉS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "5.  GESTION DES EMPLOYÉS", 1)
add_para(doc, "La page Employés (/rh/employees) centralise la gestion du personnel de votre organisation.", size=10, color=GRIS, space_after=6)

add_heading(doc, "5.1  Liste des employés", 2)
add_para(doc, "Le tableau présente les informations essentielles de chaque employé avec les colonnes suivantes :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Matricule · Nom · Prénom · Email · Téléphone · Sexe · Âge · Ancienneté")
add_bullet(doc, "Entité · Direction · Département · Fonction · Catégorie · Ville")
add_bullet(doc, "Salaire brut (visible RH/Admin/DG uniquement) · Devise · N+1 (supérieur hiérarchique)")
add_para(doc, space_after=4)

add_heading(doc, "5.2  Filtres et recherche", 2)
add_table(doc,
    ["Filtre", "Description"],
    [
        ["Recherche texte", "Recherche par nom, prénom ou matricule (insensible à la casse)"],
        ["Entité", "Filtre par entité (ECG, ELCAM, EXCA)"],
        ["Direction", "Filtre par direction"],
        ["Département", "Filtre par département"],
        ["Sexe", "M / F / Autre"],
        ["Pays / Ville", "Filtre géographique"],
        ["Tri", "Par nom, prénom, ancienneté ou âge (croissant/décroissant)"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "5.3  Actions disponibles", 2)
add_table(doc,
    ["Action", "Bouton / Icône", "Rôles autorisés", "Description"],
    [
        ["Créer un employé", "Bouton + Nouveau", "RH · Admin · DG · PCA · AG", "Ouvre le formulaire de création complet"],
        ["Modifier un employé", "Icône crayon ✏️", "RH · Admin · DG · PCA · AG", "Modifie la fiche employé existante"],
        ["Supprimer (désactiver)", "Icône poubelle 🗑️", "Admin uniquement", "Passe le statut en CONGEDIE (suppression douce)"],
        ["Import Excel/CSV", "Menu ⋯ → Importer", "Admin · RH", "Import en masse depuis fichier (voir chapitre 25)"],
        ["Export CSV", "Menu ⋯ → Exporter", "Admin · RH · DG · Directeur", "Télécharge la liste filtrée en .csv"],
        ["Télécharger modèle", "Menu ⋯ → Télécharger modèle", "Admin · RH", "Télécharge le template Excel d'import pré-rempli"],
    ],
    col_widths=[3.5, 3.5, 4, 6]
)
add_info_box(doc, "La suppression d'un employé est une suppression douce : le compte est désactivé (statut CONGEDIE) mais les données sont conservées. Une restauration est possible par l'administrateur.", bg="FFF0F0")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. FICHE EMPLOYÉ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "6.  FICHE EMPLOYÉ — CRÉATION ET MODIFICATION", 1)

add_heading(doc, "6.1  Champs du formulaire", 2)
add_table(doc,
    ["Catégorie", "Champs", "Obligatoire"],
    [
        ["Identité", "Matricule · Nom · Prénom · Date de naissance · Sexe (M/F/Autre)", "Matricule · Nom · Prénom"],
        ["Contact", "Email · Téléphone (+237...) · Contact d'urgence", "Non"],
        ["Professionnel", "Entité · Direction · Département · Fonction · Catégorie · Rôle", "Entité"],
        ["Emploi", "Statut (ACTIF/SUSPENDU/CONGEDIE) · Date d'embauche · N+1 (supérieur) · Fonction N+1", "Date d'embauche"],
        ["Localisation", "Ville · Pays", "Non"],
        ["Paie", "Salaire brut · Devise (XAF/EUR/USD/XOF)", "Non (visible RH/Admin)"],
        ["Congés", "Solde congés initial (jours)", "Non (défaut : 0)"],
        ["Famille", "État matrimonial · Nombre d'enfants", "Non"],
        ["Autre", "Diplôme · Années d'expérience · Observations", "Non"],
    ],
    col_widths=[3.2, 10.3, 3.4]
)

add_heading(doc, "6.2  Règles de validation", 2)
add_table(doc,
    ["Champ", "Règle"],
    [
        ["Matricule", "Alphanumérique uniquement (lettres, chiffres, tirets). Ex : EMP001, AG-042"],
        ["Email", "Format email valide. Doit être unique dans tout le système."],
        ["Âge", "Minimum 18 ans. Exception : Stagiaire et Apprenti (pas de minimum)."],
        ["Date embauche", "Doit être postérieure à la date de naissance."],
        ["Téléphone", "Format international recommandé : +237600000000 (Cameroun par défaut)"],
        ["Catégorie", "Valeur stricte : Cadre supérieur · Cadre moyen · Agent de maîtrise · Agent qualifié · Agent non qualifié · Apprenti · Stagiaire"],
        ["Statut", "ACTIF (défaut) · SUSPENDU · CONGEDIE"],
        ["Salaire", "Nombre positif, sans symbole monétaire. Ex : 850000"],
    ],
    col_widths=[3.5, 13.4]
)

add_heading(doc, "6.3  Photo de profil", 2)
add_para(doc, "Cliquez sur l'avatar dans la fiche employé pour téléverser une photo. Formats acceptés : JPG, PNG. Taille maximale recommandée : 2 Mo.", size=10, color=GRIS, space_after=4)

add_heading(doc, "6.4  Fonctionnalités complémentaires depuis la fiche", 2)
add_bullet(doc, "Voir la Timeline / Parcours de l'employé (historique des mutations, promotions)")
add_bullet(doc, "Générer un PDF de la fiche employé")
add_bullet(doc, "Créer un compte utilisateur pour l'employé (Admin uniquement)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. ADMINISTRATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "7.  ADMINISTRATION", 1)
add_para(doc, "La section Administration (/rh/administration) permet de gérer la structure organisationnelle de l'entreprise. Accessible aux rôles RH, PCA, AG et ADMIN.", size=10, color=GRIS, space_after=6)

for section_name, desc, actions in [
    ("7.1  Entités", "Les entités représentent les filiales ou sociétés du groupe (ECG, ELCAM, EXCA).",
     ["Créer une nouvelle entité (PCA/Admin)", "Renommer une entité", "Supprimer (si aucun employé lié)", "Visualiser les directions et départements rattachés"]),
    ("7.2  Directions", "Les directions sont les grandes divisions fonctionnelles (ex : Direction Financière).",
     ["Créer une direction et l'associer à une entité", "Renommer une direction", "Lier des départements à une direction"]),
    ("7.3  Départements", "Les départements sont les unités opérationnelles (ex : Comptabilité, RH).",
     ["Créer un département et le lier à une direction", "Renommer un département", "Associer des villes / implantations géographiques"]),
    ("7.4  Fonctions", "Les fonctions de référence servent de catalogue pour les intitulés de poste.",
     ["Créer une nouvelle fonction (libellé libre)", "Modifier ou supprimer une fonction", "Filtrer par département ou direction", "Ces fonctions alimentent les listes déroulantes de création d'employé"]),
]:
    add_heading(doc, section_name, 2)
    add_para(doc, desc, size=10, color=GRIS, space_after=4)
    for a in actions:
        add_bullet(doc, a)
    add_para(doc, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. GESTION DES UTILISATEURS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "8.  GESTION DES UTILISATEURS (ADMIN)", 1)
add_para(doc, "Accessible depuis /rh/utilisateurs. Réservé aux rôles ADMIN, PCA et AG.", size=10, color=GRIS, space_after=6)

add_heading(doc, "8.1  Tableau des utilisateurs", 2)
add_para(doc, "Le tableau affiche : Matricule · Nom · Prénom · Email · Rôle · MFA activé · Compte actif · Dernier accès.", size=10, color=GRIS, space_after=4)

add_heading(doc, "8.2  Actions disponibles", 2)
add_table(doc,
    ["Action", "Description"],
    [
        ["Modifier le rôle", "Dropdown : EMPLOYE · RESPONSABLE · DIRECTEUR · DG · RH · DFC · PCA · AG · ADMIN"],
        ["Activer / Désactiver", "Toggle ON/OFF — un compte désactivé ne peut plus se connecter"],
        ["Activer / Désactiver MFA", "Toggle pour forcer ou retirer l'authentification multi-facteurs"],
        ["Réinitialiser le mot de passe", "Génère un mot de passe temporaire et l'envoie par email à l'utilisateur"],
        ["Créer un compte pour un employé", "Associe un compte de connexion à un dossier employé existant"],
    ],
    col_widths=[5, 11.9]
)
add_info_box(doc, "Important : la création d'un employé via import ou via le formulaire NE crée PAS automatiquement un compte de connexion. L'administrateur doit créer le compte séparément depuis cette page ou depuis la fiche employé.", bg="FFF0F0")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. CONGÉS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "9.  CONGÉS", 1)
add_para(doc, "La gestion des congés est accessible depuis /rh/conges. Elle vous permet de soumettre, suivre et valider les demandes de congé.", size=10, color=GRIS, space_after=6)

add_heading(doc, "9.1  Soumettre une demande de congé", 2)
add_para(doc, "Cliquez sur le bouton ➕ Nouvelle demande. Remplissez le formulaire :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Date de début *")
add_bullet(doc, "Date de fin *")
add_bullet(doc, "Motif (texte libre, optionnel)")
add_bullet(doc, "Remplaçant (optionnel : sélectionnez un collègue)")
add_para(doc, space_after=4)
add_info_box(doc, "La durée en jours ouvrables est calculée automatiquement (lundi au vendredi, week-ends exclus). Votre solde restant est affiché en temps réel.")

add_heading(doc, "9.2  Règles importantes", 2)
add_table(doc,
    ["Règle", "Détail"],
    [
        ["Durée minimum", "1 jour ouvrable"],
        ["Délai de dépôt", "La demande doit être créée au moins 3 semaines avant la date de début. Elle ne peut plus être modifiée 2 semaines avant le départ."],
        ["Solde suffisant", "Le nombre de jours demandés ne peut pas dépasser votre solde de congés restant."],
        ["Dates cohérentes", "La date de fin doit être postérieure à la date de début."],
    ],
    col_widths=[4.5, 12.4]
)

add_heading(doc, "9.3  Suivi des demandes — Onglet Envoyé", 2)
add_para(doc, "L'onglet « Envoyé » liste toutes vos demandes de congé avec :", size=10, color=GRIS, space_after=4)
add_table(doc,
    ["Colonne", "Description"],
    [
        ["Date début / fin", "Période du congé demandé"],
        ["Durée", "Nombre de jours ouvrables calculés"],
        ["Motif", "Texte saisi à la création"],
        ["Statut", "En attente · Validé ✅ · Refusé ❌"],
        ["État workflow", "-- · AttenteRH · Active · ClotureDemandee · Cloturee"],
        ["Actions", "Voir détail · Modifier · Supprimer · Workflow · Imprimer PDF"],
    ],
    col_widths=[4, 12.9]
)
add_info_box(doc, "Vous pouvez modifier ou supprimer une demande uniquement si elle est encore en statut « En attente » et que le délai de modification n'est pas écoulé (2 semaines avant départ).")

add_heading(doc, "9.4  Valider une demande — Onglet Reçu", 2)
add_para(doc, "Si vous êtes N+1, RH ou validateur, l'onglet « Reçu » affiche les demandes en attente de votre validation :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "✅ Approuver : valide la demande (commentaire optionnel)")
add_bullet(doc, "❌ Refuser : rejette la demande (motif de refus obligatoire)")
add_bullet(doc, "Assigner un remplaçant : ouvre une modale de sélection d'employé")
add_para(doc, space_after=4)
add_info_box(doc, "Le chemin de validation standard est : Demandeur → N+1 → RH → DG/AG/PCA. Chaque étape envoie une notification automatique au validateur suivant.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. PERMISSIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "10.  PERMISSIONS", 1)
add_para(doc, "Les permissions couvrent les absences conventionnelles (mariage, naissance, décès, maladie, maternité). Accessible depuis /rh/permissions.", size=10, color=GRIS, space_after=6)

add_heading(doc, "10.1  Types de permissions disponibles", 2)
add_table(doc,
    ["Type", "Sous-type", "Durée", "Justificatif requis"],
    [
        ["Mariage", "Mariage du travailleur", "4 jours", "Acte de mariage certifié"],
        ["Mariage", "Mariage d'un enfant", "2 jours", "Acte de mariage certifié"],
        ["Paternité", "Naissance (conjoint/e)", "3 jours", "Certificat ou acte de naissance"],
        ["Baptême", "Baptême d'un enfant", "1 jour", "Attestation établissement religieux"],
        ["Décès", "Décès conjoint/e", "5 jours", "Acte de décès"],
        ["Décès", "Décès d'un enfant", "3 jours", "Acte de décès"],
        ["Décès", "Décès d'un parent", "5 jours", "Acte de décès"],
        ["Décès", "Décès beau-parent", "3 jours", "Acte de décès"],
        ["Décès", "Décès frère/sœur", "3 jours", "Acte de décès"],
        ["Maladie", "Avec justificatif médical", "3 jours", "Certificat médical"],
        ["Maternité", "Normale", "14 semaines (98 j)", "Certificat d'accouchement"],
        ["Maternité", "Pathologique", "20 semaines (140 j)", "Certificat d'accouchement"],
    ],
    col_widths=[3, 5, 3.5, 5.4]
)

add_heading(doc, "10.2  Soumettre une demande de permission", 2)
add_para(doc, "Cliquez sur ➕ Nouvelle demande et remplissez :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Type de permission (liste déroulante)")
add_bullet(doc, "Sous-type (liste imbriquée — se met à jour selon le type choisi)")
add_bullet(doc, "Date de début * (la date de fin est calculée automatiquement selon la durée légale)")
add_bullet(doc, "Motif détaillé")
add_bullet(doc, "Documents justificatifs (téléverser PDF/JPG)")
add_para(doc, space_after=4)
add_info_box(doc, "Le RH peut demander les justificatifs manquants. La permission peut être refusée si les justificatifs ne sont pas fournis dans les délais.")

add_heading(doc, "10.3  Validation et suivi", 2)
add_para(doc, "Le circuit de validation est identique aux congés (N+1 → RH → DG/AG/PCA). Les onglets « Envoyé » et « Reçu » fonctionnent de la même manière.", size=10, color=GRIS, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. DEMANDES DE SORTIES
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "11.  DEMANDES DE SORTIES PROFESSIONNELLES", 1)
add_para(doc, "Les demandes de sorties (/rh/sorties) permettent de déclarer une absence temporaire durant les heures de travail (rendez-vous médical, urgence familiale, course professionnelle).", size=10, color=GRIS, space_after=6)

add_heading(doc, "11.1  Soumettre une demande de sortie", 2)
add_bullet(doc, "Date et heure de départ *")
add_bullet(doc, "Heure de retour prévue *")
add_bullet(doc, "Motif * (texte libre)")
add_bullet(doc, "Lieu (optionnel)")
add_para(doc, space_after=4)

add_heading(doc, "11.2  Validation", 2)
add_para(doc, "La demande est transmise directement à votre N+1. Une notification en temps réel est envoyée. La validation est rapide (un seul niveau de validation).", size=10, color=GRIS, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12. MISSIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "12.  MISSIONS", 1)
add_para(doc, "Le module Missions (/rh/missions) gère les déplacements professionnels. Il supporte les missions simples, multi-destinations (segments) et multi-missionnaires.", size=10, color=GRIS, space_after=6)

add_heading(doc, "12.1  Créer une mission", 2)
add_para(doc, "Cliquez sur ➕ Nouvelle mission. Le formulaire comprend :", size=10, color=GRIS, space_after=4)
add_table(doc,
    ["Champ", "Obligatoire", "Description"],
    [
        ["Titre de la mission", "Oui", "Libellé court et descriptif"],
        ["Destination (Pays / Ville)", "Oui", "Autocomplete mondial — pays et villes du monde entier"],
        ["Missionnaires", "Oui", "Multi-sélection d'employés. Âge minimum 18 ans requis."],
        ["Date de début", "Oui", "Date de départ en mission"],
        ["Date de fin", "Oui", "Date de retour prévue (doit être > date début)"],
        ["Heure d'arrivée", "Non", "Heure d'arrivée à destination (défaut : 18h00)"],
        ["Heure de retour", "Non", "Heure de retour prévue (défaut : 17h00)"],
        ["Motif / Description", "Oui", "Objectifs et contexte de la mission"],
        ["Segments", "Non", "Ajoutez des étapes intermédiaires (voir 12.3)"],
    ],
    col_widths=[4.5, 2.2, 10.2]
)

add_heading(doc, "12.2  Règles importantes pour les missions", 2)
add_table(doc,
    ["Règle", "Détail"],
    [
        ["Âge minimum missionnaire", "18 ans stricts. Tout missionnaire mineur déclenche une erreur de validation."],
        ["Destination valide", "Le pays et la ville doivent exister dans la base de données géographique mondiale."],
        ["Multi-missionnaires", "Plusieurs employés peuvent être assignés à la même mission."],
        ["Durée de mission", "La date de fin doit être postérieure ou égale à la date de début."],
        ["Rapport obligatoire", "Un rapport de mission (PDF ou DOCX) doit être téléversé avant de pouvoir clôturer."],
        ["Frais attachés", "Les frais peuvent être soumis jusqu'à la clôture de la mission."],
    ],
    col_widths=[4.5, 12.4]
)

add_heading(doc, "12.3  Segments de mission (multi-destinations)", 2)
add_para(doc, "Pour les missions comportant plusieurs étapes géographiques, activez les segments :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Ajoutez autant de segments que nécessaire (Ville · Pays · Activité · Transport utilisé)")
add_bullet(doc, "La durée de chaque étape est calculée automatiquement")
add_bullet(doc, "Exemple : Yaoundé (J1-J2) → Douala (J3) → Libreville (J4-J5)")
add_para(doc, space_after=4)

add_heading(doc, "12.4  Suivi des missions — Onglet Envoyé", 2)
add_table(doc,
    ["Colonne", "Description"],
    [
        ["Mission", "Titre de la mission"],
        ["Destination", "Pays de destination principal"],
        ["Missionnaires", "Liste des employés en mission"],
        ["Dates", "Période de la mission"],
        ["Durée", "Nombre de jours calculé"],
        ["Segments", "Nombre d'étapes intermédiaires"],
        ["Frais", "Montant total des frais soumis"],
        ["Statut", "En attente · Validé · Refusé"],
        ["État workflow", "-- · AttenteRH · Active · ClotureDemandee · Cloturee"],
        ["Actions", "Voir · Modifier · Supprimer · Rapport · Workflow · Clôturer"],
    ],
    col_widths=[3.2, 13.7]
)

add_heading(doc, "12.5  Clôturer une mission", 2)
add_para(doc, "La clôture d'une mission indique que le missionnaire est de retour et que tout est finalisé. Elle nécessite :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Rapport de mission téléversé (PDF ou DOCX, minimum 1 Mo)")
add_bullet(doc, "Frais de mission finalisés et justificatifs téléversés")
add_bullet(doc, "Toutes les validations RH/DG obtenues")
add_para(doc, space_after=4)
add_info_box(doc, "La clôture déclenche une notification récapitulative aux validateurs et archive la mission dans l'historique.", bg="E8F5E9")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 13. FRAIS DE MISSION
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "13.  FRAIS DE MISSION", 1)
add_para(doc, "Les frais de mission (/rh/frais) sont associés à chaque mission. Chaque missionnaire peut soumettre ses frais engagés pendant la mission.", size=10, color=GRIS, space_after=6)

add_heading(doc, "13.1  Types de frais", 2)
add_table(doc,
    ["Type de frais", "Calcul", "Exemple"],
    [
        ["Frais transport", "Montant unitaire × nombre de jours/trajets", "Billet avion, taxi"],
        ["Frais hôtel", "Montant unitaire × nombre de nuits", "Chambre d'hôtel"],
        ["Frais de déplacement", "Montant fixe ou unitaire", "Location voiture"],
        ["Frais mission divers", "Montant libre", "Repas, fournitures"],
    ],
    col_widths=[4.5, 5, 7.4]
)

add_heading(doc, "13.2  Soumettre des frais", 2)
add_bullet(doc, "Sélectionnez la mission concernée")
add_bullet(doc, "Renseignez les montants pour chaque type de frais (saisie numérique, sans symbole monétaire)")
add_bullet(doc, "Téléversez les justificatifs (factures, reçus, photos) au format PDF ou JPG")
add_bullet(doc, "Cliquez sur Enregistrer")
add_para(doc, space_after=4)

add_heading(doc, "13.3  Validation et paiement", 2)
add_para(doc, "Le workflow de validation des frais suit le circuit : N+1 → RH → DG → Paiement/Comptabilité. Une notification est envoyée à chaque étape.", size=10, color=GRIS, space_after=4)
add_info_box(doc, "Les frais ne peuvent plus être modifiés ou supprimés une fois validés par un validateur.", bg="FFF0F0")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 14. OPÉRATIONS / WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "14.  OPÉRATIONS / WORKFLOW UNIFIÉ", 1)
add_para(doc, "La page Opérations (/rh/operations) est une inbox unifiée regroupant tous vos flux de travail : congés, permissions, missions, frais, sorties. C'est le tableau de bord central des validations.", size=10, color=GRIS, space_after=6)

add_heading(doc, "14.1  Deux onglets principaux", 2)
add_bullet(doc, "Envoyé : toutes vos demandes soumises (tous types confondus)")
add_bullet(doc, "Reçu : demandes en attente de votre validation (si vous êtes validateur)")
add_para(doc, space_after=4)

add_heading(doc, "14.2  Filtres disponibles", 2)
add_table(doc,
    ["Filtre", "Valeurs"],
    [
        ["Mois", "Filtre par mois de début de l'opération"],
        ["Statut", "En attente · Validé · Refusé · Tous"],
        ["Source (type)", "Congé · Permission · Mission · Frais · Sortie · Tous"],
        ["Émetteur", "Recherche libre par nom/matricule"],
        ["État workflow", "-- · AttenteRH · Active · ClotureDemandee · Cloturee"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "14.3  États du workflow", 2)
add_table(doc,
    ["État", "Signification"],
    [
        ["--", "En cours de création ou brouillon"],
        ["AttenteRH", "Opération soumise, en attente de traitement par les RH"],
        ["Active", "Approuvée et en cours d'exécution"],
        ["ClotureDemandee", "Le demandeur a signalé la fin de l'opération"],
        ["Cloturee", "Opération entièrement finalisée et archivée"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "14.4  Détail d'une opération", 2)
add_para(doc, "Cliquez sur « Voir » pour ouvrir la vue détaillée. Elle affiche :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Le résumé de l'opération (type, dates, parties, montants)")
add_bullet(doc, "La progression de validation : timeline verticale avec avatars des validateurs")
add_bullet(doc, "Les commentaires et notes à chaque étape du workflow")
add_bullet(doc, "Les boutons Valider / Refuser (avec commentaire) si vous êtes le validateur courant")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 15. REMPLAÇANTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "15.  REMPLAÇANTS", 1)
add_para(doc, "La gestion des remplaçants (/rh/remplacants) permet d'identifier la personne qui assure les responsabilités d'un employé pendant son absence.", size=10, color=GRIS, space_after=6)
add_bullet(doc, "Assigner un remplaçant : depuis l'approbation d'un congé/permission, ouvrez la modale de sélection et choisissez un employé disponible")
add_bullet(doc, "Visualiser qui remplace qui : tableau récapitulatif Absent · Période · Remplaçant · Statut")
add_bullet(doc, "Désassigner : si la situation change avant le départ")

# ═══════════════════════════════════════════════════════════════════════════
# 16. CALENDRIER DES CONGÉS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "16.  CALENDRIER DES CONGÉS", 1)
add_para(doc, "Le calendrier (/rh/calendrier-conge) offre une vue temporelle de l'ensemble des absences.", size=10, color=GRIS, space_after=6)

add_heading(doc, "16.1  Navigation et vues", 2)
add_bullet(doc, "Vue mensuelle, hebdomadaire ou annuelle (boutons de sélection)")
add_bullet(doc, "Navigation mois par mois avec les flèches de navigation")
add_bullet(doc, "Filtres : Entité · Direction · Département · Équipe · Mois/Trimestre/Année")
add_para(doc, space_after=4)

add_heading(doc, "16.2  Code couleur", 2)
add_table(doc,
    ["Couleur", "Signification"],
    [
        ["Vert 🟩", "Congé validé — absence confirmée"],
        ["Jaune 🟨", "Congé en attente de validation"],
        ["Rouge 🟥", "Congé refusé"],
        ["Gris  ⬜", "Employé absent pour autre raison (maladie, permission, etc.)"],
    ],
    col_widths=[4, 12.9]
)
add_para(doc, space_after=4)
add_info_box(doc, "Survolez ou cliquez sur une entrée du calendrier pour voir le détail (dates, motif, remplaçant) et accéder directement à la demande.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 17. NOTIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "17.  NOTIFICATIONS", 1)
add_para(doc, "Le centre de notifications (/rh/notifications) centralise tous les événements qui vous concernent. Le badge rouge sur la cloche indique le nombre de notifications non lues.", size=10, color=GRIS, space_after=6)

add_heading(doc, "17.1  Types de notifications", 2)
add_table(doc,
    ["Type", "Déclencheur", "Destinataire"],
    [
        ["VALIDATION ✅", "Votre demande a été approuvée", "Demandeur"],
        ["REFUS ❌", "Votre demande a été refusée (avec motif)", "Demandeur"],
        ["ALERTE CONGÉS 🟠", "Votre solde de congés est inférieur à 5 jours", "Employé"],
        ["RAPPEL DÉPART 🔵", "Votre mission démarre demain", "Missionnaire + N+1"],
        ["RAPPEL RETOUR 🔵", "Votre mission se termine aujourd'hui", "Missionnaire + N+1"],
        ["DEMANDE MISSION 🟣", "Vous avez été assigné(e) à une nouvelle mission", "Missionnaire"],
        ["DEMANDE EXPLICATION 🟡", "Une explication vous est demandée", "Employé"],
        ["ÉVALUATION 🔷", "Un cycle d'évaluation est ouvert / rappel de complétion", "Évalué"],
        ["CLÔTURE REQUISE 🌸", "Une opération doit être clôturée", "Responsable clôture"],
        ["PAIEMENT 💚", "Vos frais ont été approuvés et seront payés", "Missionnaire"],
        ["NOUVEL EMPLOYÉ ⚪", "Un employé vient d'être créé — créer son compte", "Admin uniquement"],
        ["IMPORT TERMINÉ ⚪", "Résumé : N ajoutés / M échecs", "Admin uniquement"],
    ],
    col_widths=[4.5, 6.5, 6]
)

add_heading(doc, "17.2  Canaux de notification", 2)
add_bullet(doc, "In-app (interface web) : visible dans le centre de notifications, badge sur la cloche")
add_bullet(doc, "Email : envoyé automatiquement sur l'adresse professionnelle enregistrée")
add_bullet(doc, "Push navigateur : si vous avez accepté les notifications du navigateur lors de la première connexion")
add_para(doc, space_after=4)

add_heading(doc, "17.3  Gérer ses notifications", 2)
add_bullet(doc, "Marquer comme lu : cliquez sur l'icône œil 👁️ ou ouvrez la notification")
add_bullet(doc, "Supprimer : cliquez sur la corbeille 🗑️")
add_bullet(doc, "Voir le détail : cliquez sur le lien dans la notification pour accéder directement à l'opération concernée")
add_bullet(doc, "Filtrer par type : utilisez le dropdown en haut de la liste")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 18. ÉVALUATIONS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "18.  ÉVALUATIONS", 1)
add_para(doc, "Le module Évaluations (/rh/evaluations) gère les cycles d'appréciation du personnel.", size=10, color=GRIS, space_after=6)

add_heading(doc, "18.1  Processus d'évaluation", 2)
add_para(doc, "Un cycle d'évaluation se déroule en trois étapes :", size=10, color=GRIS, space_after=4)
add_bullet(doc, "Étape 1 — Auto-évaluation : l'employé remplit sa propre grille d'appréciation")
add_bullet(doc, "Étape 2 — Évaluation manager : le N+1 évalue l'employé en croisant avec l'auto-évaluation")
add_bullet(doc, "Étape 3 — Validation RH : les RH valident et archivent les résultats")
add_para(doc, space_after=4)

add_heading(doc, "18.2  Grille de notation", 2)
add_para(doc, "Les compétences, comportements et objectifs sont notés selon l'échelle : Insuffisant · Satisfaisant · Bon · Très bon · Excellent (1 à 5).", size=10, color=GRIS, space_after=4)

add_heading(doc, "18.3  Actions disponibles", 2)
add_bullet(doc, "Créer un nouveau cycle d'évaluation (RH/Admin)")
add_bullet(doc, "Remplir une évaluation (bouton « Évaluer »)")
add_bullet(doc, "Consulter les résultats et synthèses")
add_bullet(doc, "Générer le PDF de synthèse")
add_bullet(doc, "Recevoir des rappels automatiques si l'évaluation n'est pas complétée")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 19. FICHES DE POSTE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "19.  FICHES DE POSTE", 1)
add_para(doc, "Les fiches de poste (/rh/fiche-de-poste) décrivent les fonctions de chaque emploi de l'organisation.", size=10, color=GRIS, space_after=6)

add_table(doc,
    ["Section de la fiche", "Contenu"],
    [
        ["Titre du poste / Fonction", "Intitulé officiel du poste"],
        ["Description générale", "Contexte et raison d'être du poste"],
        ["Missions principales", "Liste des responsabilités et activités clés"],
        ["Compétences requises", "Savoir-faire techniques et comportementaux"],
        ["Rattachement hiérarchique", "Direction, département, N+1 direct"],
        ["Conditions d'emploi", "Catégorie, statut, localisation, télétravail"],
    ],
    col_widths=[5, 11.9]
)
add_para(doc, space_after=4)
add_bullet(doc, "Toutes les fiches sont consultables par tous les employés (lecture seule)")
add_bullet(doc, "La création et modification sont réservées aux RH")
add_bullet(doc, "Export PDF disponible depuis chaque fiche")

# ═══════════════════════════════════════════════════════════════════════════
# 20. ORGANIGRAMME
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "20.  ORGANIGRAMME INTERACTIF", 1)
add_para(doc, "L'organigramme (/rh/orgchart) visualise la hiérarchie complète de l'organisation sous forme de graphique interactif.", size=10, color=GRIS, space_after=6)

add_heading(doc, "20.1  Navigation dans l'organigramme", 2)
add_bullet(doc, "Zoom avant/arrière avec la molette de la souris ou les boutons + / -")
add_bullet(doc, "Déplacement par glisser-déposer (pan) de la vue")
add_bullet(doc, "Cliquez sur une carte employé pour voir son profil complet")
add_bullet(doc, "Filtrez par entité ou direction via les sélecteurs en haut de page")
add_para(doc, space_after=4)

add_heading(doc, "20.2  Informations affichées sur chaque carte", 2)
add_bullet(doc, "Photo de profil · Nom et prénom · Fonction / Poste")
add_bullet(doc, "Email et téléphone (si droits suffisants)")
add_bullet(doc, "Entité / Direction / Département")
add_para(doc, space_after=4)

add_heading(doc, "20.3  Règle visuelle", 2)
add_info_box(doc, "Le supérieur hiérarchique (N+1) est toujours positionné au-dessus de l'employé dans l'arborescence. Si un employé apparaît au même niveau qu'un responsable, cela indique que le lien hiérarchique (champ N+1) n'est pas renseigné dans sa fiche.", bg="FFF0F0")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 21. ORGANISATION & LOCALISATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "21.  ORGANISATION & LOCALISATION", 1)
add_para(doc, "Ce module (/rh/organisation) permet de gérer la structure géographique et organisationnelle : pays, villes, entités implantées par localisation.", size=10, color=GRIS, space_after=6)

add_heading(doc, "21.1  Onglets disponibles", 2)
add_table(doc,
    ["Onglet", "Description", "Actions (Admin/PCA)"],
    [
        ["Pays", "Liste des pays avec codes ISO où le groupe est implanté", "Ajouter · Renommer · Supprimer"],
        ["Villes", "Villes par pays", "Ajouter · Renommer · Supprimer"],
        ["Entités (par localisation)", "Arborescence Pays → Ville → Entités implantées", "Créer · Lier à une ville"],
        ["Directions (par localisation)", "Vue : Pays → Ville → Entité → Directions", "Visualisation et gestion"],
        ["Départements (par localisation)", "Zoom direction → depts opérationnels", "Visualisation et gestion"],
    ],
    col_widths=[3.5, 6.5, 6.9]
)
add_para(doc, space_after=4)
add_info_box(doc, "La recherche géographique mondiale est disponible : tapez les premières lettres d'un pays ou d'une ville pour obtenir une liste de suggestions (autocomplete mondial).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 22. PROFIL PERSONNEL ET SÉCURITÉ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "22.  PROFIL PERSONNEL ET SÉCURITÉ", 1)
add_para(doc, "Votre profil (/rh/profile) centralise vos informations personnelles et vos paramètres de sécurité.", size=10, color=GRIS, space_after=6)

add_heading(doc, "22.1  Informations personnelles", 2)
add_bullet(doc, "Identité : nom, prénom, photo de profil")
add_bullet(doc, "Coordonnées : email, téléphone, contact d'urgence")
add_bullet(doc, "Informations professionnelles : matricule, entité, direction, département, fonction, catégorie, rôle, N+1")
add_bullet(doc, "Statut d'emploi et dates (embauche, naissance)")
add_para(doc, space_after=4)
add_info_box(doc, "Vous pouvez modifier vos informations personnelles (email, téléphone, photo). Le matricule, la direction, le département et le rôle ne peuvent être modifiés que par un administrateur ou RH.")

add_heading(doc, "22.2  Sécurité du compte", 2)
add_table(doc,
    ["Action", "Description"],
    [
        ["Activer le MFA", "Scannez le QR code avec votre application d'authentification (Google Authenticator, Authy)"],
        ["Changer de mot de passe", "Saisissez l'ancien puis le nouveau (minimum 8 caractères, lettres + chiffres)"],
        ["Sessions actives", "Voir toutes les sessions ouvertes et les déconnecter à distance"],
        ["Historique de connexion", "Liste des dernières connexions (date, heure, IP, navigateur)"],
    ],
    col_widths=[4.5, 12.4]
)

add_heading(doc, "22.3  Préférences", 2)
add_bullet(doc, "Thème visuel : clair ou sombre (bascule immédiate)")
add_bullet(doc, "Notifications push : activer ou désactiver les notifications du navigateur")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 23. ANALYTICS RH AVANCÉ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "23.  ANALYTICS RH AVANCÉ", 1)
add_para(doc, "Accessible aux rôles RH, DG, PCA et ADMIN depuis /rh/analytics. Ce module offre des tableaux de bord analytiques avancés sur le capital humain.", size=10, color=GRIS, space_after=6)

add_table(doc,
    ["Dashboard", "Métriques clés"],
    [
        ["Effectifs", "Évolution headcount · Pyramide des âges · Distribution par catégorie · Turnover (entrées/sorties)"],
        ["Congés", "Jours utilisés/restants par employé · Prévisions vs réalisé · Tendances saisonnières"],
        ["Missions", "Nombre par destination · Coût moyen · Durée moyenne · Taux de clôture"],
        ["Paie", "Masse salariale par entité/direction · Distribution salaires · Évolution annuelle"],
    ],
    col_widths=[3.5, 13.4]
)

# ═══════════════════════════════════════════════════════════════════════════
# 24. PARCOURS EMPLOYÉ (TIMELINE)
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "24.  PARCOURS EMPLOYÉ (TIMELINE)", 1)
add_para(doc, "La timeline (/rh/timeline) retrace l'historique complet de carrière d'un employé sous forme de frise chronologique.", size=10, color=GRIS, space_after=6)

add_table(doc,
    ["Type d'événement", "Description"],
    [
        ["EMBAUCHE", "Date d'entrée dans l'organisation et poste initial"],
        ["PROMOTION", "Changement de grade ou de catégorie"],
        ["MUTATION", "Transfert vers une autre direction / entité"],
        ["TRANSFERT", "Changement d'entité (inter-filiale)"],
        ["CONGEDIE", "Fin du contrat (date et motif)"],
        ["AUTRE", "Tout autre événement marquant (formation diplômante, etc.)"],
    ],
    col_widths=[4, 12.9]
)
add_para(doc, space_after=4)
add_bullet(doc, "Les RH peuvent créer, modifier et supprimer des événements de carrière")
add_bullet(doc, "Un PDF de l'historique complet peut être généré depuis la timeline")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 25. IMPORT & EXPORT DES EMPLOYÉS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "25.  IMPORT & EXPORT DES EMPLOYÉS", 1)

add_heading(doc, "25.1  Importer des employés en masse", 2)
add_para(doc, "L'import permet de créer plusieurs employés en une seule opération depuis un fichier Excel ou CSV.", size=10, color=GRIS, space_after=6)

add_heading(doc, "Procédure d'import", 3)
add_para(doc, "1.  Téléchargez le modèle Excel : page Employés → menu ⋯ → Télécharger modèle Excel.", size=10, color=GRIS, space_after=3)
add_para(doc, "2.  Ouvrez le fichier TEMPLATE_IMPORT_EMPLOYES_EMS.xlsx dans Excel ou LibreOffice.", size=10, color=GRIS, space_after=3)
add_para(doc, "3.  Effacez les lignes d'exemple grisées (lignes 4 et 5 de l'onglet Employés).", size=10, color=GRIS, space_after=3)
add_para(doc, "4.  Remplissez vos données, une ligne par employé. Utilisez les listes déroulantes disponibles.", size=10, color=GRIS, space_after=3)
add_para(doc, "5.  Enregistrez le fichier en .xlsx.", size=10, color=GRIS, space_after=3)
add_para(doc, "6.  Dans EMS : page Employés → menu ⋯ → Importer → sélectionnez votre fichier.", size=10, color=GRIS, space_after=3)
add_para(doc, "7.  Un résumé s'affiche : nombre d'employés importés avec succès et erreurs (numéro de ligne + motif).", size=10, color=GRIS, space_after=6)

add_heading(doc, "Structure du fichier modèle", 3)
add_table(doc,
    ["Onglet", "Contenu"],
    [
        ["Employés", "La feuille à remplir (25 colonnes, 300 lignes pré-équipées, 2 lignes d'exemple à effacer)"],
        ["Instructions", "Procédure d'import, règles métier, format des dates, valeurs acceptées"],
        ["Référence", "Valeurs réelles de la BD (entités, directions, depts, fonctions, rôles, villes)"],
    ],
    col_widths=[3.5, 13.4]
)
add_para(doc, space_after=4)

add_heading(doc, "Champs obligatoires", 3)
add_table(doc,
    ["Champ (colonne Excel)", "Règle"],
    [
        ["matricule", "Unique, alphanumérique. Ex : EMP001"],
        ["nom", "Nom de famille"],
        ["prenom", "Prénom"],
        ["date_embauche", "Format : aaaa-mm-jj  Ex : 2024-01-15"],
        ["entite", "Valeur exacte : ECG · ELCAM · EXCA"],
    ],
    col_widths=[5, 11.9]
)
add_para(doc, space_after=4)
add_info_box(doc, "Les administrateurs reçoivent une notification pour chaque employé importé, ainsi qu'un résumé global à la fin de l'import. La création des comptes utilisateurs est une étape distincte réalisée par l'administrateur.", bg="E8F5E9")

add_heading(doc, "25.2  Exporter la liste des employés", 2)
add_para(doc, "Page Employés → menu ⋯ → Exporter en CSV. Le fichier généré contient tous les employés selon vos filtres actifs, avec toutes les colonnes (hors salaire pour les rôles non autorisés).", size=10, color=GRIS, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 26. RÈGLES MÉTIER CLÉS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "26.  RÈGLES MÉTIER CLÉS — RÉCAPITULATIF", 1)

add_heading(doc, "26.1  Employés", 2)
add_table(doc,
    ["Règle", "Détail"],
    [
        ["Matricule", "Alphanumérique + tirets uniquement. Ex : EMP001, AG-042. Unique en base de données."],
        ["Âge minimum", "18 ans. Exception : Stagiaire et Apprenti (pas de minimum d'âge)."],
        ["Email", "Format RFC valide. Unique en base de données. Optionnel à la création."],
        ["Catégorie", "Enum strict : Cadre supérieur · Cadre moyen · Agent de maîtrise · Agent qualifié · Agent non qualifié · Apprenti · Stagiaire"],
        ["Statut emploi", "ACTIF (défaut) · SUSPENDU · CONGEDIE"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "26.2  Congés", 2)
add_table(doc,
    ["Règle", "Détail"],
    [
        ["Durée minimum", "1 jour ouvrable"],
        ["Délai de modification", "Modification impossible 2 semaines avant la date de début"],
        ["Délai de dépôt", "Demande à créer au moins 3 semaines avant le départ"],
        ["Solde", "Le nombre de jours demandés ne peut pas dépasser le solde disponible"],
        ["Cohérence dates", "Date de fin strictement postérieure à la date de début"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "26.3  Missions", 2)
add_table(doc,
    ["Règle", "Détail"],
    [
        ["Âge missionnaire", "18 ans stricts pour tout missionnaire assigné"],
        ["Rapport avant clôture", "Rapport de mission obligatoire (PDF/DOCX) pour pouvoir clôturer"],
        ["Frais avant clôture", "Tous les frais doivent être soumis et validés avant clôture"],
        ["Destination", "Pays et ville doivent exister dans la base géographique mondiale"],
    ],
    col_widths=[4, 12.9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 27. SÉCURITÉ ET CONTRÔLE D'ACCÈS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "27.  SÉCURITÉ ET CONTRÔLE D'ACCÈS", 1)

add_heading(doc, "27.1  Authentification", 2)
add_table(doc,
    ["Mécanisme", "Détail"],
    [
        ["JWT (Token)", "Chaque session génère un token sécurisé valable 1 jour"],
        ["MFA (TOTP)", "Second facteur optionnel via application d'authentification (Google Authenticator, Authy)"],
        ["Sessions trackées", "Chaque connexion est enregistrée (date, heure, adresse IP, navigateur)"],
        ["Déconnexion auto", "La session expire après une période d'inactivité configurable par l'admin"],
    ],
    col_widths=[4, 12.9]
)

add_heading(doc, "27.2  Contrôle d'accès (RBAC)", 2)
add_table(doc,
    ["Niveau de contrôle", "Description"],
    [
        ["Route (page)", "Chaque page vérifie le rôle avant affichage — redirection si accès refusé"],
        ["Interface (boutons)", "Les boutons sensibles sont masqués ou désactivés selon le rôle"],
        ["Données (API)", "Chaque endpoint API filtre les données selon le rôle et le périmètre de l'utilisateur"],
    ],
    col_widths=[4.5, 12.4]
)

add_heading(doc, "27.3  Bonnes pratiques recommandées", 2)
add_bullet(doc, "Activez le MFA sur votre compte, en particulier pour les rôles Admin, PCA, AG, DG et RH")
add_bullet(doc, "Ne partagez jamais votre mot de passe avec un collègue")
add_bullet(doc, "Déconnectez-vous en quittant votre poste de travail, surtout sur un ordinateur partagé")
add_bullet(doc, "Si vous suspectez un accès non autorisé, changez immédiatement votre mot de passe et prévenez l'administrateur")
add_bullet(doc, "Signalez toute anomalie ou comportement suspect à l'équipe SI")
add_para(doc, space_after=6)

add_separator(doc)
add_para(doc, "MANUEL D'UTILISATION EMS — Version 1.0 — Avril 2026  ·  Elite Capital Group", 
         size=8, color=GRIS, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Document confidentiel à usage interne exclusivement — Toute reproduction non autorisée est interdite.",
         size=8, color=GRIS, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Sauvegarde ────────────────────────────────────────────────────────────
out_path = (
    r"c:\Users\cedri\OneDrive - ELITE CAPITAL Group S.A"
    r"\Documents\EMS\extranet\MANUEL_UTILISATION_EMS.docx"
)
doc.save(out_path)
print(f"Saved → {out_path}")
print("Manuel d'utilisation EMS généré avec succès.")
print("  Police : Century Gothic")
print("  Charte EMS : #02162E · #D0202B · #606060")
print("  Sections : 27 chapitres complets")
